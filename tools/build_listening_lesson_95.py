from __future__ import annotations

import argparse
import os
import re
import shutil
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / ".tmp" / "python-packages"))

import eng_to_ipa as ipa  # type: ignore
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Data" / "listening" / "are_you_a_workaholic"
AUDIO_DIR = OUT_DIR / "audio"
TMP_DIR = OUT_DIR / "_tmp"

TITLE = "ARE YOU A WORKAHOLIC? (BẠN CÓ PHẢI LÀ NGƯỜI NGHIỆN CÔNG VIỆC?) - CEFR: B2"

LESSON = [
    (
        "In a culture that encourages ambition, to ensure a stable income, workers often devote all their time, including time for self-care, to handle an enormous workload.",
        "Trong một nền văn hóa khuyến khích sự cầu tiến, để đảm bảo nguồn thu nhập ổn định, người lao động thường dành hết thời gian, bao gồm cả thời gian chăm sóc bản thân, để xử lý khối lượng công việc khổng lồ.",
    ),
    (
        "This phenomenon is particularly noticeable among the younger generation.",
        "Hiện tượng này đặc biệt được nhận thấy rõ hơn ở giới trẻ.",
    ),
    (
        'Many young people today prefer to be seen as "workaholics" because the feeling of being busy makes them feel like they are creating value for themselves and society.',
        'Nhiều bạn trẻ hiện nay thích được xem là những "workaholic" (người nghiện công việc) bởi cảm giác bận rộn khiến họ thấy mình đang tạo ra giá trị cho bản thân và xã hội.',
    ),
    (
        "However, it’s time for young workers to pay attention to the warnings about their physical and mental health and reconsider their relationship with work.",
        "Thế nhưng đã đến lúc người lao động trẻ cần chú ý đến những báo động về thể chất lẫn tinh thần, đồng thời nhìn nhận lại mối quan hệ giữa mình và công việc.",
    ),
    (
        '"Workaholism" was first used in 1971 by psychologist Wayne Oates.',
        'Cụm từ "nghiện công việc" lần đầu tiên được sử dụng vào năm 1971 bởi nhà tâm lý học Wayne Oates.',
    ),
    (
        "He defined it as a term describing a person with an uncontrollable need to work.",
        "Ông định nghĩa đây là cụm từ mô tả tình trạng một người có nhu cầu làm việc ngoài kiểm soát.",
    ),
    (
        "Workaholics always think of ways to spend as much time as possible on work.",
        'Những "workaholic" luôn nghĩ cách để dành thời gian cho công việc nhiều nhất có thể.',
    ),
    (
        "They are constantly obsessed and feel guilty whenever they leave the workspace or take time to play or rest.",
        "Họ luôn bị ám ảnh và mang cảm giác tội lỗi mỗi khi rời khỏi không gian làm việc hay dành thời gian vui chơi, nghỉ dưỡng.",
    ),
    (
        "Workaholism is becoming increasingly common among young people, especially those who are perfectionists and pursue perfectionism.",
        "Hội chứng nghiện công việc hiện nay ngày càng phổ biến hơn ở giới trẻ, đặc biệt là những người cầu toàn, theo đuổi chủ nghĩa hoàn hảo.",
    ),
    (
        "The feeling of physical, emotional, and mental exhaustion due to work has become all too normal for young people.",
        "Cảm giác kiệt quệ về thể chất, cảm xúc và tinh thần vì công việc là điều đã trở nên quá đỗi bình thường đối với những người trẻ tuổi.",
    ),
    (
        "We ignore the fact that we borrow hours of sleep to work extra hours and even forget about self-care.",
        "Chúng ta phớt lờ sự thật rằng chúng ta mượn số giờ ngủ để làm thêm giờ, và thậm chí quên đi việc chăm sóc bản thân.",
    ),
    (
        "This phenomenon leads to the consequence that nearly half of millennials have left their jobs due to mental health reasons, and diagnoses of depression are increasing at a faster rate among the younger generation than any other age group.",
        "Hiện tượng này kéo theo hệ quả là gần một nửa số người thuộc thế hệ thiên niên kỷ đã rời bỏ công việc vì lý do sức khỏe tâm thần, và các chẩn đoán trầm cảm đang gia tăng với tốc độ nhanh hơn ở thế hệ thanh niên so với bất kỳ nhóm tuổi nào khác.",
    ),
    (
        "Taking on too many tasks simultaneously with high frequency, and prioritizing work above all else are signs of a workaholic.",
        "Ôm đồm quá nhiều việc cùng lúc với tần suất cao, coi công việc là ưu tiên hàng đầu... đó là dấu hiệu của người nghiện việc.",
    ),
    (
        "Maintaining this habit causes many to fall into a state of stress and anxiety.",
        "Duy trì thói quen này khiến nhiều người rơi vào trạng thái căng thẳng, lo âu.",
    ),
    (
        "Even if they aren’t working, they feel uncomfortable.",
        "Thậm chí nếu không làm việc, họ sẽ có cảm giác khó chịu.",
    ),
    (
        "Those who suffer from workaholism are at risk of health problems such as insomnia, fatigue, stomach issues, anxiety, and depression.",
        "Người mắc chứng nghiện việc có nguy cơ gặp vấn đề về sức khỏe như mất ngủ, mệt mỏi, vấn đề về dạ dày, lo âu và trầm cảm.",
    ),
    (
        "The lack of balance between work and personal life also leads to conflicts in relationships.",
        "Sự thiếu cân bằng giữa công việc và cuộc sống cá nhân cũng dẫn đến mâu thuẫn trong các mối quan hệ.",
    ),
    (
        "Many cases result in having no time for oneself, family, children, or social relationships, gradually leading to loneliness.",
        "Nhiều trường hợp không có thời gian cho bản thân, gia đình, con cái, các mối quan hệ xã hội, dần dần thành cô đơn.",
    ),
    (
        'When they stop working, they feel alone and empty, so they continue to "work addictively" to relieve this, creating a vicious cycle.',
        'Khi ngừng làm việc, họ cảm thấy đơn độc, trống rỗng, nên lại tiếp tục "nghiện việc" để giải tỏa, tạo thành vòng luẩn quẩn.',
    ),
    (
        "What solutions are there for workaholics?",
        'Giải pháp nào cho những "workaholic"?',
    ),
    (
        "Workaholism is entirely treatable, as long as you recognize your problem.",
        "Chứng nghiện việc hoàn toàn có thể điều trị được, miễn là bạn nhận ra vấn đề của mình.",
    ),
    (
        "Here are some methods you can implement to overcome workaholism and achieve a more balanced life.",
        "Sau đây là một số phương pháp bạn có thể thực hiện để thoát khỏi chứng nghiện việc và có một cuộc sống cân bằng hơn.",
    ),
    (
        "Understand what is important.",
        "Hiểu rõ điều gì là quan trọng.",
    ),
    (
        "Thoroughly research your to-do list and major goals.",
        "Nghiên cứu thật kỹ danh sách những việc phải làm và những mục tiêu lớn của bạn.",
    ),
    (
        "Some tasks are essential, and you must complete them yourself, but most others are not.",
        "Một số việc là trọng yếu và bạn phải tự thực hiện, nhưng hầu hết những việc còn lại thì không.",
    ),
    (
        'Ask yourself: "How will this help my company?" and determine to eliminate unnecessary things.',
        'Tự hỏi mình: "Việc này sẽ giúp gì cho công ty của tôi?" và hãy quyết tâm gạt bỏ những thứ không cần thiết.',
    ),
    (
        "Reorganize your daily activities.",
        "Sắp xếp lại hoạt động trong ngày.",
    ),
    (
        "Working smart means you must maximize the time when you can work most effectively.",
        "Làm việc thông minh đồng nghĩa với việc bạn phải tận dụng tối đa khoảng thời gian bạn có thể làm việc hiệu quả nhất.",
    ),
    (
        "Mark this time on your schedule for the most creative and important tasks you need to do.",
        "Hãy đánh dấu trên thời gian biểu của mình để dành khoảng thời gian đó cho những việc sáng tạo và quan trọng nhất bạn cần làm.",
    ),
    (
        "Don’t get swept up in other people's schedules or demands.",
        "Đừng bị cuốn theo lịch trình hay nhu cầu của người khác.",
    ),
    (
        "Working smart allows you and your mind to rest so that you can return to work with a refreshed and clearer mindset.",
        "Làm việc thông minh tạo điều kiện cho bạn và trí não bạn nghỉ ngơi, để sau đó có thể quay lại công việc với một tâm trí tươi mới và minh mẫn hơn.",
    ),
    (
        "As a result, you will handle everything more efficiently during your work hours.",
        "Nhờ đó, trong khoảng thời gian làm việc bạn sẽ xử lý mọi vấn đề hiệu quả hơn.",
    ),
    (
        "You can plan a walk, write a journal, play sports, or have dinner with people after work.",
        "Bạn có thể lên kế hoạch đi dạo, viết nhật ký, chơi thể thao hoặc ăn tối với mọi người sau giờ làm việc.",
    ),
    (
        "Creating a new routine will help you forget your obsession with working at the office.",
        "Việc tạo ra một thói quen mới sẽ giúp bạn quên đi sự ám ảnh của công việc ở cơ quan.",
    ),
    (
        "The important thing is to find a hobby that suits you.",
        "Điều quan trọng là bạn cần tìm thấy sở thích phù hợp với bản thân.",
    ),
    (
        "When you engage in these activities, you will be distracted and no longer think about work.",
        "Khi tham gia những hoạt động đó, bạn sẽ bị phân tâm, không còn nhớ tới công việc nữa.",
    ),
    (
        "Sacrificing health and joy is not a smart way to pursue the long career path ahead.",
        "Đánh đổi sức khỏe và niềm vui không phải là cách thông minh để theo đuổi con đường sự nghiệp còn dài phía trước.",
    ),
    (
        "Successful people know their time is valuable, but they do not spend it all on work; they also participate in outside activities.",
        "Những người thành công biết rằng thời gian của họ là quý giá nhưng họ không hoàn toàn dành cho công việc mà còn tham gia những hoạt động bên ngoài.",
    ),
    (
        "A balance between work and life will help each person be happier and more energized, leading to higher efficiency at work.",
        "Sự cân bằng giữa công việc và cuộc sống sẽ giúp mỗi người hạnh phúc hơn, tràn đầy năng lượng để đạt được hiệu quả cao hơn trong công việc.",
    ),
]


IPA_OVERRIDES = {
    "workaholism": "ˌwɝkəˈhɔˌlɪzəm",
    "workaholics": "ˌwɝkəˈhɔlɪks",
    "workaholic": "ˌwɝkəˈhɔlɪk",
    "wayne": "weɪn",
    "oates": "oʊts",
    "millennials": "mɪˈlɛniəlz",
    "to-do": "tə ˈdu",
    "self-care": "ˌsɛlf ˈkɛr",
    "workspace": "ˈwɝkˌspeɪs",
    "addictively": "əˈdɪktɪvli",
}


def clean_ipa(text: str) -> str:
    normalized = text.replace("’", "'").replace("“", '"').replace("”", '"')
    converted = ipa.convert(normalized)
    for word, value in IPA_OVERRIDES.items():
        converted = re.sub(rf"\b{re.escape(word)}\*?\b", value, converted, flags=re.IGNORECASE)
        converted = re.sub(re.escape(word), value, converted, flags=re.IGNORECASE)
    converted = converted.replace("1971", "ˌnaɪnˈtin ˌsɛvənˈti wən")
    converted = converted.replace('"', "")
    converted = converted.replace("*", "")
    converted = converted.replace("ˈɪts", "ɪts")
    converted = re.sub(r"\s+", " ", converted).strip()
    return f"/{converted}/"


def set_run_font(run, name: str, size: int | None = None, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def build_text(rows: list[tuple[str, str, str]]) -> str:
    blocks = [TITLE, ""]
    for english, vietnamese, phonetic in rows:
        blocks.extend([english, vietnamese, phonetic, ""])
    return "\n".join(blocks).rstrip() + "\n"


def build_docx(rows: list[tuple[str, str, str]], output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(TITLE)
    set_run_font(title_run, "Arial", 18, True)

    for index, (english, vietnamese, phonetic) in enumerate(rows, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{index:02d}. {english}")
        set_run_font(r, "Arial", 11, True)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(vietnamese)
        set_run_font(r, "Arial", 11)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(phonetic)
        set_run_font(r, "Arial", 10)
        r.italic = True

    doc.save(output_path)


def generate_gtts_audio(rows: list[tuple[str, str, str]], audio_dir: Path) -> None:
    from gtts import gTTS  # type: ignore

    audio_dir.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)

    for index, (english, _, _) in enumerate(rows, start=1):
        filename = f"95_{index:02d}_are_you_a_workaholic.mp3"
        path = audio_dir / filename
        if path.exists() and path.stat().st_size > 0:
            continue
        tts = gTTS(text=english, lang="en", tld="com", slow=False)
        tts.save(path)


def generate_openai_audio(rows: list[tuple[str, str, str]], audio_dir: Path, voice: str) -> None:
    from openai import OpenAI

    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError(
            "OPENAI_API_KEY is not set. Export it first, then rerun this script."
        )

    audio_dir.mkdir(parents=True, exist_ok=True)
    client = OpenAI()

    for index, (english, _, _) in enumerate(rows, start=1):
        filename = f"95_{index:02d}_are_you_a_workaholic.mp3"
        path = audio_dir / filename
        if path.exists() and path.stat().st_size > 0:
            continue

        with client.audio.speech.with_streaming_response.create(
            model="gpt-4o-mini-tts",
            voice=voice,
            input=english,
            instructions=(
                "Speak clearly in natural American English at a moderate pace "
                "for English dictation practice. Keep the delivery warm, calm, "
                "and human, with crisp consonants and natural sentence rhythm."
            ),
            response_format="mp3",
        ) as response:
            response.stream_to_file(path)


def write_manifest(rows: list[tuple[str, str, str]], audio_dir_name: str, output_path: Path) -> None:
    manifest = [
        "lesson_id,title,cefr,index,audio_file,english,vietnamese,ipa",
    ]
    for index, (english, vietnamese, phonetic) in enumerate(rows, start=1):
        audio_file = f"{audio_dir_name}/95_{index:02d}_are_you_a_workaholic.mp3"
        values = ["95", "Are You a Workaholic?", "B2", str(index), audio_file, english, vietnamese, phonetic]
        manifest.append(",".join('"' + value.replace('"', '""') + '"' for value in values))
    output_path.write_text("\n".join(manifest) + "\n", encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--tts-provider", choices=["gtts", "openai"], default="gtts")
    parser.add_argument("--voice", default="marin")
    parser.add_argument("--skip-doc", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = [(english, vietnamese, clean_ipa(english)) for english, vietnamese in LESSON]

    if not args.skip_doc:
        (OUT_DIR / "are_you_a_workaholic.txt").write_text(build_text(rows), encoding="utf-8")
        build_docx(rows, OUT_DIR / "are_you_a_workaholic.docx")

    if args.tts_provider == "openai":
        audio_dir_name = f"audio_openai_{args.voice}"
        generate_openai_audio(rows, OUT_DIR / audio_dir_name, args.voice)
        write_manifest(rows, audio_dir_name, OUT_DIR / f"manifest_openai_{args.voice}.csv")
    else:
        generate_gtts_audio(rows, AUDIO_DIR)
        write_manifest(rows, "audio", OUT_DIR / "manifest.csv")

    if TMP_DIR.exists():
        shutil.rmtree(TMP_DIR)

    print(f"Wrote {OUT_DIR}")
    print(f"Sentences: {len(rows)}")


if __name__ == "__main__":
    main()
